import os
import uuid
import PyPDF2
from docx import Document
from docx.shared import Pt, Inches
from fastapi import UploadFile
import shutil
import logging
from typing import  Tuple

from app.core.ai_processor import AIProcessor
from app.core.ragflow import upload_files_to_dataset

# 创建日志记录器
logger = logging.getLogger("pdf_handler")


class PDFHandler:
    """PDF处理工具类"""
    
    def __init__(self, upload_dir: str, output_dir: str, base_url: str, ai_base_url: str = None, ai_model: str = None):
        self.upload_dir = upload_dir
        self.output_dir = output_dir
        self.base_url = base_url
        os.makedirs(upload_dir, exist_ok=True)
        os.makedirs(output_dir, exist_ok=True)
        
        # 初始化AI处理器
        self.ai_processor = None
        if ai_base_url and ai_model:
            self.ai_processor = AIProcessor(ai_base_url, ai_model)
            logger.info(f"AI处理器已初始化: {ai_base_url}, 模型: {ai_model}")
    
    async def save_pdf(self, file: UploadFile) -> str:
        """
        保存上传的PDF文件
        
        Args:
            file: 上传的PDF文件
            
        Returns:
            保存的PDF文件路径
        """
        # 生成唯一文件名
        unique_id = str(uuid.uuid4())
        base_name = os.path.splitext(file.filename)[0]
        pdf_filename = f"{base_name}_{unique_id}.pdf"
        pdf_path = os.path.join(self.upload_dir, pdf_filename)
        
        # 保存文件
        with open(pdf_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        file_size = os.path.getsize(pdf_path)
        logger.info(f"保存PDF文件: {pdf_path}, 大小: {file_size} 字节")
        
        return pdf_path
    
    def extract_pdf_text(self, pdf_path: str) -> str:
        """
        从PDF文件中提取文本
        
        Args:
            pdf_path: PDF文件路径
            
        Returns:
            提取的文本内容
        """
        logger.info(f"开始从PDF提取文本: {pdf_path}")
        
        try:
            # 用于存储所有提取的文本
            all_text = []
            
            # 打开PDF文件
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # 获取PDF页数
                total_pages = len(reader.pages)
                logger.info(f"PDF共有 {total_pages} 页")
                
                # 处理每一页
                for page_num in range(total_pages):
                    # 获取页面
                    page = reader.pages[page_num]
                    
                    # 提取文本
                    text = page.extract_text()
                    
                    # 添加到文本集合
                    if text:
                        all_text.append(f"--- 第 {page_num + 1} 页 ---\n{text}")
                    
                    logger.info(f"已提取第 {page_num + 1} 页文本")
            
            # 合并所有文本
            full_text = "\n\n".join(all_text)
            logger.info(f"文本提取完成，总长度: {len(full_text)}")
            
            return full_text
            
        except Exception as e:
            logger.error(f"提取PDF文本时出错: {str(e)}")
            return f"提取文本时出错: {str(e)}"
    
    def convert_pdf_to_word(self, pdf_path: str,enter_text:str=None) -> Tuple[str, int, str, str, str]:
        """
        将PDF文件转换为Word文档
        
        Args:
            pdf_path: PDF文件路径
            enter_text: 额外的文本输入
        
        Returns:
            tuple: (输出文件的路径, 总页数, 提取的文本内容, 处理后的文本内容, markdown文件路径)
        """
        # 生成输出文件名
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        docx_filename = f"{base_name}.docx"
        output_path = os.path.join(self.output_dir, docx_filename)
        
        # 生成Markdown文件名
        markdown_filename = f"{base_name}.md"
        markdown_path = os.path.join(self.output_dir, markdown_filename)
        
        logger.info(f"开始转换PDF到Word: {pdf_path} -> {output_path}")
        
        try:
            # 创建一个新的Word文档
            doc = Document()
            
            # 设置页边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # 提取文本
            full_text = self.extract_pdf_text(pdf_path)
            
            # 打开PDF文件获取页数
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                total_pages = len(reader.pages)
                
                # 处理每一页
                for page_num in range(total_pages):
                    # 获取页面
                    page = reader.pages[page_num]
                    
                    # 提取文本
                    text = page.extract_text()
                    
                    # 添加页码标题
                    doc.add_heading(f'第 {page_num + 1} 页', level=1)
                    
                    # 添加提取的文本
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(text)
                    run.font.size = Pt(11)  # 设置字体大小
                    
                    # 添加分页符（除了最后一页）
                    if page_num < total_pages - 1:
                        doc.add_page_break()
                    
                    logger.info(f"已处理第 {page_num + 1} 页")
            
            # 使用AI处理文本
            processed_text = ""
            if self.ai_processor and full_text:
                logger.info("使用AI处理提取的文本")
                processed_text = self.ai_processor.process_text(full_text,enter_text=enter_text)
                
                # 保存为Markdown文件
                with open(markdown_path, "w", encoding="utf-8") as md_file:
                    md_file.write(processed_text)
                logger.info(f"已保存Markdown文件: {markdown_path}")
                # 将markdown文件上传至ragflow
                upload_files_to_dataset([markdown_path])
            else:
                logger.info("未配置AI处理器或文本为空，跳过AI处理")
            
            # 保存Word文档
            doc.save(output_path)
            logger.info(f"转换完成！已保存为: {output_path}")
            
            return output_path, total_pages, full_text, processed_text, markdown_path
            
        except Exception as e:
            logger.error(f"转换PDF到Word时出错: {str(e)}")
            # 创建一个包含错误信息的文档
            error_doc = Document()
            error_doc.add_heading("转换错误", level=1)
            error_doc.add_paragraph(f"转换PDF文件时出现错误: {str(e)}")
            error_doc.save(output_path)
            
            # 返回错误文档、0页和空文本
            return output_path, 0, "", "", "" 